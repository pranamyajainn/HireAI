import os
import re
import json
import logging
from typing import Dict, List, Any
from datetime import datetime
import fitz  # PyMuPDF for better OCR
import docx
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self, gemini_api_key=None):
        """Initialize OCR + LLM Resume Parser using Gemini"""
        self.gemini_api_key = gemini_api_key or os.getenv('GEMINI_API_KEY')
        
        # Initialize Gemini client
        if self.gemini_api_key:
            try:
                genai.configure(api_key=self.gemini_api_key)
                self.model = genai.GenerativeModel('gemini-2.0-flash')
                self.ai_available = True
                logger.info("✅ Gemini AI client initialized successfully")
            except Exception as e:
                logger.error(f"❌ Gemini AI initialization failed: {e}")
                self.model = None
                self.ai_available = False
        else:
            logger.warning("⚠️ No Gemini API key provided - using basic parsing")
            self.model = None
            self.ai_available = False    
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        """Main function to parse a resume file using OCR + AI"""
        try:
            logger.info(f"🔍 Extracting text from: {os.path.basename(file_path)}")
            
            # Extract text using advanced OCR
            text = self._extract_text_ocr(file_path)
            
            if not text or len(text.strip()) < 50:
                return {"error": "Could not extract sufficient text from file"}
            
            logger.info(f"📄 Extracted {len(text)} characters")
            
            # Parse using AI if available, otherwise use basic parsing
            if self.ai_available:
                parsed_data = self._parse_with_ai(text)
            else:
                parsed_data = self._parse_text_basic(text)
            
            # Add metadata
            parsed_data['raw_text'] = text
            parsed_data['file_path'] = file_path
            parsed_data['parsed_at'] = datetime.now().isoformat()
            parsed_data['parser_type'] = 'OCR+AI' if self.ai_available else 'OCR+Basic'
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"❌ Error parsing resume: {e}")
            return {"error": f"Resume parsing failed: {str(e)}"}
    
    # --- ENHANCED OCR DEBUGGING AND EXTRACTION ---
    def debug_ocr_extraction(self, text: str):
        """Prints a summary of what was extracted from OCR text for debugging."""
        logger.debug("\n===== OCR Extraction Debug =====")
        logger.debug(f"Length: {len(text)} characters")
        logger.debug(f"First 500 chars:\n{text[:500]}")
        logger.debug(f"Last 300 chars:\n{text[-300:]}")
        logger.debug("===== END OCR Debug =====\n")    
    def _extract_from_pdf_ocr(self, file_path: str) -> str:
        """Extract text from PDF using advanced straight-line extraction for better resume parsing.
        
        This method uses multiple extraction strategies:
        1. Extract text in reading order with directional awareness
        2. Extract text as blocks with spatial recognition
        3. Extract text as structured HTML to preserve layouts
        4. Use raw text extraction to capture text that might be missed
        5. Combine all strategies for maximum coverage
        """
        text = ""
        doc = None
        try:
            doc = fitz.open(file_path)
            logger.info(f"📑 PDF has {len(doc)} pages")
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = ""
                
                # STRATEGY 1: Directional extraction (best for columnar text)
                # This preserves the reading order in columns and tables
                directional_text = page.get_text("text", sort=True)
                
                # STRATEGY 2: Block extraction (best for grouped elements)
                blocks_text = ""
                blocks = page.get_text("blocks")
                # Sort blocks by vertical position (top to bottom)
                sorted_blocks = sorted(blocks, key=lambda b: b[1])  # Sort by y0 coordinate
                for block in sorted_blocks:
                    if block[6] == 0:  # Text block (not image)
                        blocks_text += block[4] + "\n"
                
                # STRATEGY 3: Raw text extraction (catches text missed by other methods)
                raw_text = page.get_text("text", sort=False)
                
                # STRATEGY 4: HTML extraction (preserves some formatting)
                html_text = ""
                try:
                    html = page.get_text("html")
                    # Basic HTML cleaning to extract just text
                    html_text = re.sub(r'<[^>]+>', ' ', html)
                    html_text = re.sub(r'\s+', ' ', html_text).strip()
                except:
                    pass
                
                # Combine strategies with weights (prefer directional for resumes)
                combined_texts = []
                if directional_text.strip():
                    combined_texts.append(directional_text)
                if blocks_text.strip():
                    combined_texts.append(blocks_text)
                if raw_text.strip() and len(raw_text) > len(directional_text) * 1.2:  # Only if raw adds 20% more content
                    combined_texts.append(raw_text)
                if html_text.strip() and len(html_text) > len(directional_text) * 1.1:  # Only if HTML adds 10% more content
                    combined_texts.append(html_text)
                
                # If all extraction methods failed, try image-based OCR as last resort
                if not combined_texts:
                    try:
                        import importlib.util
                        has_pil = importlib.util.find_spec("PIL") is not None
                        has_pytesseract = importlib.util.find_spec("pytesseract") is not None
                        
                        if has_pil and has_pytesseract:
                            from PIL import Image
                            import pytesseract
                            # Higher resolution matrix for better OCR
                            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))  # 3x resolution for better OCR
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            # Try multiple OCR configurations for best results
                            page_text = pytesseract.image_to_string(
                                img, 
                                config='--psm 1 --oem 3'  # PSM 1 = Auto page segmentation with OSD
                            )
                            logger.info(f"🔍 Used image-based OCR for page {page_num+1}")
                    except Exception as e:
                        logger.warning(f"⚠️ Image OCR failed: {e}")
                
                # Join all text with best strategy first
                page_text = "\n".join(combined_texts)
                
                # Post-process to clean up text
                page_text = re.sub(r'(\n\s*){3,}', '\n\n', page_text)  # Remove excessive newlines
                
                text += page_text + "\n\n"
                logger.info(f"📄 Extracted page {page_num+1} with {len(page_text)} characters")
            
            # Final cleanup
            text = re.sub(r'(\n\s*){3,}', '\n\n', text)  # Remove excessive newlines again
            
            return text.strip()
        except Exception as e:
            logger.error(f"❌ PDF OCR error: {e}")
            return ""        
        finally:
            if doc is not None:
                doc.close()
                
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file, including headers, footers, and tables."""
        text = ""
        try:
            doc = docx.Document(file_path)
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            # Try to extract headers/footers if possible
            if hasattr(doc, 'sections'):
                for section in doc.sections:
                    if hasattr(section, 'header'):
                        text += section.header.text + "\n"
                    if hasattr(section, 'footer'):
                        text += section.footer.text + "\n"
        except Exception as e:
            logger.error(f"❌ DOCX extraction error: {e}")
        return text.strip()
    
    def _extract_text_ocr(self, file_path: str) -> str:
        """Extract text using advanced OCR (PDF/DOCX) and print debug info."""
        file_extension = os.path.splitext(file_path)[1].lower()
        try:
            logger.info(f"\n🔬 Extracting text from {file_extension} file: {os.path.basename(file_path)}")
            if file_extension == '.pdf':
                text = self._extract_from_pdf_ocr(file_path)
            elif file_extension in ['.docx', '.doc']:
                text = self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Print debug information
            self.debug_ocr_extraction(text)
            
            # Show detailed visualization of extracted content
            self.visualize_ocr_results(text, file_path)
            
            return text
        except Exception as e:
            logger.error(f"❌ OCR extraction error: {e}")
            return ""
    def _parse_with_ai(self, text: str) -> Dict[str, Any]:
        """Parse resume text using Gemini AI"""
        try:
            logger.info("🤖 Using AI to parse resume...")
            
            # Create a comprehensive prompt for resume parsing
            prompt = f"""
You are an expert resume parser. Parse the following resume text and extract structured information.
The text was extracted from a resume using OCR, so there might be some formatting issues or errors.
Carefully analyze the text and extract as much information as possible, even if it's not perfectly formatted.

Return ONLY a valid JSON object with the following structure:

{{
    "name": "Full name of the candidate",
    "email": "Email address",
    "phone": "Phone number",
    "location": "City, State/Country",
    "skills": ["skill1", "skill2", "skill3"],
    "experience": [
        {{
            "title": "Job title",
            "company": "Company name",
            "duration": "Duration (e.g., '2020-2023')",
            "description": "Brief description of role"
        }}
    ],
    "education": [
        {{
            "degree": "Degree type",
            "institution": "University/College name",
            "year": "Graduation year or period",
            "field": "Field of study"
        }}
    ],
    "experience_years": 5,
    "summary": "Professional summary in 1-2 sentences",
    "certifications": ["cert1", "cert2"],
    "languages": ["language1", "language2"],
    "projects": [
        {{
            "name": "Project name",
            "description": "Brief description",
            "technologies": ["tech1", "tech2"]
        }}
    ]
}}

IMPORTANT: Be thorough and extract ALL relevant information. Look for patterns that might indicate:
- Names at the top of the resume
- Contact information including email addresses, phone numbers, and LinkedIn profiles
- Skills sections, usually containing comma or bullet-separated keywords
- Experience sections with company names, job titles, dates, and descriptions
- Education sections with degrees, institutions, and dates
- Project sections with project names, descriptions, and technologies used

Resume text to parse:
{text}

Return only the JSON object, no other text or explanations.
"""
            
            # Call Gemini API
            response = self.model.generate_content(prompt)
            ai_response = response.text.strip()

            # Clean the response (remove markdown code blocks if present)
            if ai_response.startswith('```json'):
                ai_response = ai_response[7:]
            if ai_response.startswith('```'):
                ai_response = ai_response[3:]
            if ai_response.endswith('```'):
                ai_response = ai_response[:-3]
            
            # Parse JSON
            try:
                parsed_data = json.loads(ai_response.strip())
                logger.info("✅ AI parsing completed successfully")
                return parsed_data
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON parsing error: {e}")
                logger.error(f"Raw AI response: {ai_response[:200]}...")
                return self._parse_text_basic(text)
        except Exception as e:
            logger.error(f"❌ AI parsing error: {e}")
            return self._parse_text_basic(text)
    
    def _parse_text_basic(self, text: str) -> Dict[str, Any]:
        """Basic parsing fallback when AI is not available"""
        logger.info("🔧 Using basic text parsing...")
        
        parsed_data = {
            'name': self._extract_name_basic(text),
            'email': self._extract_email_basic(text),
            'phone': self._extract_phone_basic(text),
            'location': self._extract_location_basic(text),
            'skills': self._extract_skills_basic(text),
            'experience': self._extract_experience_basic(text),
            'education': self._extract_education_basic(text),
            'experience_years': self._calculate_experience_years_basic(text),
            'summary': self._extract_summary_basic(text),
            'certifications': [],
            'languages': [],
            'projects': []
        }
        
        return parsed_data
    
    def _extract_name_basic(self, text: str) -> str:
        """Extract candidate name from resume text"""
        lines = text.strip().split('\n')
        
        # Usually the name is in the first few lines
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 2 and len(line.split()) <= 4:
                # Skip common headers
                skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'profile', 'contact', 'email', 'phone']
                if not any(word in line.lower() for word in skip_words):
                    # Check if it looks like a name (contains letters, minimal numbers)
                    if re.match(r'^[A-Za-z\s\.-]+$', line) and len(line) > 5:
                        return line.title()
        
        return "Name not found"
    
    def _extract_email_basic(self, text: str) -> str:
        """Extract email address from resume text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else "Email not found"
    
    def _extract_phone_basic(self, text: str) -> str:
        """Extract phone number from resume text"""
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{1,3}[-.\s]?\d{10}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        
        return "Phone not found"
    
    def _extract_location_basic(self, text: str) -> str:
        """Extract location/address from resume text"""
        # Look for common location patterns
        location_patterns = [
            r'[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}',  # City, ST 12345
            r'[A-Za-z\s]+,\s*[A-Z]{2}',  # City, ST
            r'[A-Za-z\s]+,\s*[A-Za-z\s]+',  # City, Country/State
        ]
        
        for pattern in location_patterns:
            locations = re.findall(pattern, text)
            if locations:
                return locations[0]
        
        return "Location not found"
    
    def _extract_skills_basic(self, text: str) -> List[str]:
        """Extract skills from resume text"""
        # Common technical skills (expanded list)
        common_skills = [
            'python', 'java', 'javascript', 'react', 'angular', 'vue', 'node.js', 'express',
            'django', 'flask', 'fastapi', 'sql', 'mysql', 'postgresql', 'mongodb', 'redis',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github',
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
            'data analysis', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'tableau',
            'html', 'css', 'bootstrap', 'sass', 'typescript', 'jquery', 'webpack',
            'microservices', 'rest api', 'graphql', 'agile', 'scrum', 'devops',
            'linux', 'windows', 'macos', 'bash', 'powershell', 'c++', 'c#', 'go',
            'rust', 'swift', 'kotlin', 'php', 'ruby', 'scala', 'r', 'matlab',
            'figma', 'sketch', 'photoshop', 'illustrator', 'ui/ux', 'design',
            'jira', 'confluence', 'slack', 'teams', 'communication', 'leadership'
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in common_skills:
            if skill in text_lower:
                found_skills.append(skill.title())
        
        # Remove duplicates and return
        return list(set(found_skills))
    
    def _extract_experience_basic(self, text: str) -> List[Dict[str, str]]:
        """Extract work experience from resume text"""
        experience = []
        
        # Look for common experience section headers
        experience_sections = re.split(r'(?i)(experience|work\s+experience|employment|professional\s+experience)', text)
        
        if len(experience_sections) > 1:
            exp_text = experience_sections[-1][:1000]  # Take first 1000 chars after experience header
            
            # Look for job titles and companies (simplified pattern)
            job_patterns = [
                r'([A-Za-z\s]+(?:Engineer|Developer|Manager|Analyst|Specialist|Coordinator|Director|Lead))[^\n]*\n[^\n]*([A-Za-z\s&,\.]+)(?:Inc|LLC|Corp|Company|Ltd)?',
                r'([A-Z][a-z]+\s+[A-Z][a-z]+)[^\n]*\n[^\n]*([A-Z][a-z\s&,\.]+)'
            ]
            
            for pattern in job_patterns:
                matches = re.findall(pattern, exp_text)
                for match in matches[:3]:  # Limit to first 3 matches
                    experience.append({
                        'title': match[0].strip(),
                        'company': match[1].strip(),
                        'duration': 'Not specified',
                        'description': 'Details extracted from resume'
                    })
                if matches:
                    break
        
        return experience
    
    def _extract_education_basic(self, text: str) -> List[Dict[str, str]]:
        """Extract education information from resume text"""
        education = []
        
        # Look for common degree patterns
        degree_patterns = [
            r'(Bachelor|Master|PhD|B\.S\.|M\.S\.|B\.A\.|M\.A\.|MBA)[^\n]*([A-Za-z\s]+(?:University|College|Institute))',
            r'(B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|PhD)[^\n]*([A-Za-z\s]+(?:University|College|Institute))'
        ]
        
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                education.append({
                    'degree': match[0].strip(),
                    'institution': match[1].strip(),
                    'year': 'Not specified',
                    'field': 'Not specified'
                })
        
        return education
    
    def _calculate_experience_years_basic(self, text: str) -> int:
        """Calculate total years of experience"""
        # Look for experience mentions
        exp_patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'(\d+)\+?\s*yrs?\s*(?:of\s*)?experience',
            r'experience[:\s]*(\d+)\+?\s*years?'
        ]
        
        years = []
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            years.extend([int(match) for match in matches])
        
        return max(years) if years else 0
    
    def _extract_summary_basic(self, text: str) -> str:
        """Extract professional summary or objective"""
        # Look for summary/objective sections
        summary_patterns = [
            r'(?i)(summary|objective|profile)[:\s]*([^\n]*(?:\n[^\n]*){0,3})',
            r'(?i)(about\s+me|professional\s+summary)[:\s]*([^\n]*(?:\n[^\n]*){0,3})'
        ]
        
        for pattern in summary_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0][1].strip()[:200]  # First 200 chars
        
        # If no summary section, take first paragraph as summary
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 50]
        if paragraphs:
            return paragraphs[0][:200]
        
        return "No summary available"
    
    def visualize_ocr_results(self, text: str, file_path: str) -> None:
        """Print a detailed visualization of OCR results to help debug extraction issues"""
        logger.info("\n" + "="*80)
        logger.info(f"📋 OCR EXTRACTION RESULTS FOR: {os.path.basename(file_path)}")
        logger.info("="*80)
        
        # Basic stats
        lines = text.split('\n')
        words = text.split()
        
        logger.info(f"📊 STATISTICS:")
        logger.info(f"   - Characters: {len(text)}")
        logger.info(f"   - Lines: {len(lines)}")
        logger.info(f"   - Words: {len(words)}")
        logger.info(f"   - Avg chars per line: {len(text)/len(lines) if lines else 0:.1f}")
        
        # Content preview
        logger.info("\n📑 CONTENT PREVIEW:")
        logger.info(f"   - First 10 lines:")
        for i, line in enumerate(lines[:10]):
            if line.strip():
                logger.info(f"     {i+1}: {line[:80]}{'...' if len(line) > 80 else ''}")
            else:
                logger.info(f"     {i+1}: [empty line]")
        
        # Show key information detection
        logger.info("\n🔍 KEY INFORMATION DETECTION:")
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        logger.info(f"   - Emails found: {', '.join(emails[:3]) if emails else 'None'}")
        
        phone_pattern = r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        logger.info(f"   - Phones found: {', '.join(phones[:3]) if phones else 'None'}")
        
        # Try to find potential name (first non-empty line that's not common header text)
        potential_name = "Not detected"
        for line in lines[:10]:
            line = line.strip()
            if len(line) > 2 and len(line.split()) <= 4:
                skip_words = ['resume', 'cv', 'curriculum', 'vitae']
                if not any(word in line.lower() for word in skip_words):
                    potential_name = line
                    break
        logger.info(f"   - Potential name: {potential_name}")
        
        # Potential skills section
        logger.info(f"   - Sections detected: {self._detect_resume_sections(text)}")
        
        logger.info("="*80 + "\n")

    def _detect_resume_sections(self, text: str) -> List[str]:
        """Detect which sections are present in the resume"""
        sections = []
        section_keywords = {
            "education": ["education", "academic", "degree", "university", "college", "school"],
            "experience": ["experience", "employment", "work history", "career"],
            "skills": ["skills", "technical skills", "competencies", "technologies"],
            "projects": ["projects", "portfolio", "works"],
            "certifications": ["certifications", "certificates", "qualifications"],
            "summary": ["summary", "profile", "objective", "about me"]
        }
        
        text_lower = text.lower()
        for section, keywords in section_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                sections.append(section)
                
        return sections